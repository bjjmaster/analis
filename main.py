from docx import Document

def finf_index(l): # долгая функция нужно будет оптимизировать
    l_of_index = []
    for z in l:
        for g in range(len(document.paragraphs)):
            if document.paragraphs[g].text == z:
                l_of_index.append(g)
    return l_of_index

def corpus(l1, l2):
    dict_of_procet = {}
    for dict in range(len(l)):
        try:
            dict_of_procet.setdefault(l1[dict], text[l2[dict] + 1 : l2[dict + 1]])
        except:
            dict_of_procet.setdefault(l1[dict], text[l2[dict] + 1: l2[dict] + 15])
    return dict_of_procet


document = Document("C:\\Users\\max\\PycharmProjects\\analis\\doc_analizator\\TZ\\good\\GOOD_Техническое задание6.docx")
text = []
for paragraph in document.paragraphs:
    text.append(paragraph.text)

l = ["Требования к подсистеме идентификации и аутентификации", "Требования к подсистеме регистрации и учёта", "Требования к подсистеме обеспечения целостности",
    "Требования к подсистеме антивирусной защиты", "Требования к подсистеме обеспечения сетевой безопасности", 'Требования к подсистеме контроля защищенности',
     'Требования к подсистеме криптографической защиты']

list_of_index = finf_index(l)
a = corpus(l, list_of_index)
for i in a:
    print(i, a[i])
